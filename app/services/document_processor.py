"""Document processor for native text extraction from various file formats."""

import asyncio
import hashlib
from typing import Optional, Dict, Any, List
from dataclasses import dataclass
import structlog

logger = structlog.get_logger(__name__)


@dataclass
class TextExtractionResult:
    """Result of text extraction from document."""
    success: bool
    text: str
    pages: int
    processing_time: float
    metadata: Optional[Dict[str, Any]] = None
    error_message: Optional[str] = None


class DocumentProcessor:
    """Process documents to extract native text content."""
    
    def __init__(self):
        """Initialize document processor."""
        self.supported_mimetypes = {
            "application/pdf": self._extract_pdf_text,
            "application/msword": self._extract_doc_text,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": self._extract_docx_text,
            "text/plain": self._extract_plain_text
        }
    
    async def extract_text(self, 
                          content: bytes, 
                          mimetype: str, 
                          timeout: int = 30) -> TextExtractionResult:
        """Extract text from document content."""
        start_time = asyncio.get_event_loop().time()
        
        try:
            # Check if mimetype is supported
            if mimetype not in self.supported_mimetypes:
                return TextExtractionResult(
                    success=False,
                    text="",
                    pages=0,
                    processing_time=asyncio.get_event_loop().time() - start_time,
                    error_message=f"Unsupported mimetype: {mimetype}"
                )
            
            # Extract text with timeout
            extractor = self.supported_mimetypes[mimetype]
            result = await asyncio.wait_for(
                extractor(content),
                timeout=timeout
            )
            
            processing_time = asyncio.get_event_loop().time() - start_time
            
            return TextExtractionResult(
                success=result["success"],
                text=result["text"],
                pages=result["pages"],
                processing_time=processing_time,
                metadata=result.get("metadata")
            )
            
        except asyncio.TimeoutError:
            processing_time = asyncio.get_event_loop().time() - start_time
            logger.error("Text extraction timed out", timeout=timeout)
            return TextExtractionResult(
                success=False,
                text="",
                pages=0,
                processing_time=processing_time,
                error_message=f"Text extraction timed out after {timeout} seconds"
            )
        except Exception as exc:
            processing_time = asyncio.get_event_loop().time() - start_time
            logger.error("Text extraction failed", exc_info=exc)
            return TextExtractionResult(
                success=False,
                text="",
                pages=0,
                processing_time=processing_time,
                error_message=str(exc)
            )
    
    async def _extract_pdf_text(self, content: bytes) -> Dict[str, Any]:
        """Extract text from PDF document."""
        try:
            # Try PyMuPDF first (faster and more reliable)
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(stream=content, filetype="pdf")
                text = ""
                pages = len(doc)
                
                for page_num in range(pages):
                    page = doc.load_page(page_num)
                    page_text = page.get_text()
                    text += page_text + "\n"
                
                doc.close()
                
                return {
                    "success": True,
                    "text": text.strip(),
                    "pages": pages,
                    "metadata": {"extractor": "pymupdf", "has_text": bool(text.strip())}
                }
                
            except ImportError:
                # Fallback to pdfplumber
                try:
                    import pdfplumber
                    import io
                    
                    with pdfplumber.open(io.BytesIO(content)) as pdf:
                        text = ""
                        pages = len(pdf.pages)
                        
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                        
                        return {
                            "success": True,
                            "text": text.strip(),
                            "pages": pages,
                            "metadata": {"extractor": "pdfplumber", "has_text": bool(text.strip())}
                        }
                        
                except ImportError:
                    # Final fallback - no PDF libraries available
                    logger.warning("No PDF libraries available for text extraction")
                    return {
                        "success": False,
                        "text": "",
                        "pages": 0,
                        "metadata": {"extractor": "none", "error": "No PDF libraries available"}
                    }
                    
        except Exception as exc:
            logger.error("PDF text extraction failed", exc_info=exc)
            return {
                "success": False,
                "text": "",
                "pages": 0,
                "metadata": {"extractor": "error", "error": str(exc)}
            }
    
    async def _extract_docx_text(self, content: bytes) -> Dict[str, Any]:
        """Extract text from DOCX document."""
        try:
            import docx
            import io
            
            doc = docx.Document(io.BytesIO(content))
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            
            return {
                "success": True,
                "text": text.strip(),
                "pages": 1,  # DOCX doesn't have explicit page count
                "metadata": {"extractor": "python-docx", "has_text": bool(text.strip())}
            }
            
        except ImportError:
            logger.warning("python-docx not available for DOCX text extraction")
            return {
                "success": False,
                "text": "",
                "pages": 0,
                "metadata": {"extractor": "none", "error": "python-docx not available"}
            }
        except Exception as exc:
            logger.error("DOCX text extraction failed", exc_info=exc)
            return {
                "success": False,
                "text": "",
                "pages": 0,
                "metadata": {"extractor": "error", "error": str(exc)}
            }
    
    async def _extract_doc_text(self, content: bytes) -> Dict[str, Any]:
        """Extract text from DOC document."""
        try:
            # Try using antiword (external tool)
            import subprocess
            import tempfile
            import os
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_file:
                temp_file.write(content)
                temp_file.flush()
                
                try:
                    result = subprocess.run(
                        ["antiword", temp_file.name],
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    
                    if result.returncode == 0:
                        text = result.stdout.strip()
                        return {
                            "success": True,
                            "text": text,
                            "pages": 1,  # DOC doesn't have explicit page count
                            "metadata": {"extractor": "antiword", "has_text": bool(text)}
                        }
                    else:
                        logger.warning("antiword failed", returncode=result.returncode, stderr=result.stderr)
                        return {
                            "success": False,
                            "text": "",
                            "pages": 0,
                            "metadata": {"extractor": "antiword", "error": result.stderr}
                        }
                        
                except subprocess.TimeoutExpired:
                    logger.warning("antiword timed out")
                    return {
                        "success": False,
                        "text": "",
                        "pages": 0,
                        "metadata": {"extractor": "antiword", "error": "Timeout"}
                    }
                finally:
                    # Clean up temp file
                    try:
                        os.unlink(temp_file.name)
                    except OSError:
                        pass
                        
        except FileNotFoundError:
            logger.warning("antiword not available for DOC text extraction")
            return {
                "success": False,
                "text": "",
                "pages": 0,
                "metadata": {"extractor": "none", "error": "antiword not available"}
            }
        except Exception as exc:
            logger.error("DOC text extraction failed", exc_info=exc)
            return {
                "success": False,
                "text": "",
                "pages": 0,
                "metadata": {"extractor": "error", "error": str(exc)}
            }
    
    async def _extract_plain_text(self, content: bytes) -> Dict[str, Any]:
        """Extract text from plain text document."""
        try:
            # Try to decode as UTF-8 first
            try:
                text = content.decode('utf-8')
            except UnicodeDecodeError:
                # Fallback to other encodings
                for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        text = content.decode(encoding)
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    # Final fallback - replace errors
                    text = content.decode('utf-8', errors='replace')
            
            return {
                "success": True,
                "text": text.strip(),
                "pages": 1,
                "metadata": {"extractor": "native", "encoding": "detected", "has_text": bool(text.strip())}
            }
            
        except Exception as exc:
            logger.error("Plain text extraction failed", exc_info=exc)
            return {
                "success": False,
                "text": "",
                "pages": 0,
                "metadata": {"extractor": "error", "error": str(exc)}
            }
    
    def supports_mimetype(self, mimetype: str) -> bool:
        """Check if processor supports the given mimetype."""
        return mimetype in self.supported_mimetypes
    
    def get_supported_mimetypes(self) -> List[str]:
        """Get list of supported mimetypes."""
        return list(self.supported_mimetypes.keys())
    
    def needs_ocr(self, result: TextExtractionResult) -> bool:
        """Determine if document needs OCR processing."""
        if not result.success:
            return True
        
        # Check if extracted text is meaningful
        text = result.text.strip()
        if not text:
            return True
        
        # Check if text is too short (might be just metadata)
        if len(text) < 50:
            return True
        
        # Check if text contains mostly whitespace or special characters
        meaningful_chars = sum(1 for c in text if c.isalnum() or c.isspace())
        if meaningful_chars / len(text) < 0.7:
            return True
        
        return False


# Global document processor instance
_document_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    """Get global document processor instance."""
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor
