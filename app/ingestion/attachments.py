"""Attachment processor for email ingestion pipeline."""

import asyncio
import time
from typing import Optional, Dict, Any, List
import structlog
import hashlib
from datetime import datetime

from .models import AttachmentInfo, OCRTask, ProcessingStatus, ErrorType
from ..services.document_processor import get_document_processor, TextExtractionResult
from ..services.ocr import get_ocr_service, OCRResult, OCRTask as OCRTaskModel
from ..config.ocr import get_ocr_settings

logger = structlog.get_logger(__name__)


class AttachmentProcessor:
    """Processes email attachments for text extraction and OCR."""
    
    def __init__(self):
        """Initialize attachment processor."""
        self.ocr_service = get_ocr_service()
        self.document_processor = get_document_processor()
        self.ocr_settings = get_ocr_settings()
        
        # OCR tasks queue (in-process for Phase-1)
        self.ocr_queue: List[OCRTaskModel] = []
        self.processing_tasks: Dict[str, OCRTaskModel] = {}
    
    async def process_attachment(self, attachment: AttachmentInfo, tenant_id: str) -> Dict[str, Any]:
        """Process attachment with text extraction and OCR if needed."""
        try:
            # Validate attachment
            if not self._is_valid_attachment(attachment):
                return {
                    "success": False,
                    "text": None,
                    "needs_ocr": False,
                    "error": "Invalid attachment",
                    "quarantined": True
                }
            
            # Try native text extraction first
            text_result = await self.document_processor.extract_text(
                attachment.content, 
                attachment.content_type,
                timeout=self.ocr_settings.default_timeout_seconds
            )
            
            if text_result.success and not self.document_processor.needs_ocr(text_result):
                # Native text extraction successful
                logger.info("Native text extracted successfully",
                           filename=attachment.filename,
                           mimetype=attachment.content_type,
                           text_length=len(text_result.text))
                
                return {
                    "success": True,
                    "text": text_result.text,
                    "needs_ocr": False,
                    "pages": text_result.pages,
                    "extractor": text_result.metadata.get("extractor", "unknown")
                }
            else:
                # Need OCR processing
                logger.info("Attachment requires OCR processing",
                           filename=attachment.filename,
                           mimetype=attachment.content_type,
                           reason="No native text or insufficient content")
                
                # Create OCR task
                ocr_task = await self._create_ocr_task(attachment, tenant_id)
                
                return {
                    "success": False,
                    "text": None,
                    "needs_ocr": True,
                    "ocr_task_id": ocr_task.task_id,
                    "pages": text_result.pages if text_result.success else 1
                }
                
        except Exception as exc:
            logger.error("Failed to process attachment",
                        filename=attachment.filename,
                        mimetype=attachment.content_type,
                        exc_info=exc)
            return {
                "success": False,
                "text": None,
                "needs_ocr": False,
                "error": str(exc),
                "quarantined": True
            }
    
    async def _create_ocr_task(self, attachment: AttachmentInfo, tenant_id: str) -> OCRTaskModel:
        """Create OCR task for attachment."""
        import time
        
        task_id = f"ocr_{attachment.filename}_{int(time.time())}"
        
        ocr_task = OCRTaskModel(
            task_id=task_id,
            attachment_id=attachment.filename,  # Using filename as ID for now
            tenant_id=tenant_id,
            content_hash=attachment.content_hash,
            mimetype=attachment.content_type,
            file_size=attachment.content_length
        )
        
        # Add to queue
        self.ocr_queue.append(ocr_task)
        logger.info("OCR task created and queued",
                   task_id=task_id,
                   filename=attachment.filename,
                   queue_size=len(self.ocr_queue))
        
        return ocr_task
    
    async def process_ocr_queue(self) -> List[Dict[str, Any]]:
        """Process queued OCR tasks."""
        results = []
        
        while self.ocr_queue and len(self.processing_tasks) < self.ocr_settings.max_concurrent_tasks:
            task = self.ocr_queue.pop(0)
            
            try:
                # Mark as processing
                task.status = "processing"
                task.started_at = time.time()
                self.processing_tasks[task.task_id] = task
                
                # Process OCR
                result = await self._process_ocr_task(task)
                results.append(result)
                
                # Remove from processing
                del self.processing_tasks[task.task_id]
                
            except Exception as exc:
                logger.error("OCR task processing failed",
                           task_id=task.task_id,
                           exc_info=exc)
                task.status = "failed"
                task.error_message = str(exc)
                results.append({
                    "task_id": task.task_id,
                    "success": False,
                    "error": str(exc)
                })
                
                # Remove from processing
                del self.processing_tasks[task.task_id]
        
        return results
    
    async def _process_ocr_task(self, task: OCRTaskModel) -> Dict[str, Any]:
        """Process individual OCR task."""
        try:
            # Get attachment content (in real implementation, fetch from storage)
            # For now, we'll simulate this
            attachment_content = b"sample content"  # Placeholder
            
            # Process OCR
            ocr_result = await self.ocr_service.extract_text(
                attachment_content,
                task.mimetype,
                language_hint=task.language_hint,
                timeout=self.ocr_settings.default_timeout_seconds
            )
            
            if ocr_result.success:
                task.status = "completed"
                task.completed_at = time.time()
                task.result = ocr_result
                
                logger.info("OCR task completed successfully",
                           task_id=task.task_id,
                           text_length=len(ocr_result.text),
                           processing_time=ocr_result.processing_time)
                
                return {
                    "task_id": task.task_id,
                    "success": True,
                    "text": ocr_result.text,
                    "confidence": ocr_result.confidence,
                    "processing_time": ocr_result.processing_time,
                    "backend": ocr_result.backend
                }
            else:
                task.status = "failed"
                task.error_message = ocr_result.error_message
                
                logger.error("OCR task failed",
                           task_id=task.task_id,
                           error=ocr_result.error_message)
                
                return {
                    "task_id": task.task_id,
                    "success": False,
                    "error": ocr_result.error_message
                }
                
        except Exception as exc:
            task.status = "failed"
            task.error_message = str(exc)
            
            logger.error("OCR task processing failed",
                       task_id=task.task_id,
                       exc_info=exc)
            
            return {
                "task_id": task.task_id,
                "success": False,
                "error": str(exc)
            }
    
    def _is_valid_attachment(self, attachment: AttachmentInfo) -> bool:
        """Check if attachment is valid for processing."""
        # Check mimetype allowlist
        if attachment.content_type not in self.ocr_settings.allowed_mimetypes:
            logger.warning("Attachment mimetype not allowed",
                          filename=attachment.filename,
                          mimetype=attachment.content_type)
            return False
        
        # Check size limits
        max_size_mb = self.ocr_settings.size_limits_mb.get(attachment.content_type, 10)
        if attachment.content_length > max_size_mb * 1024 * 1024:
            logger.warning("Attachment exceeds size limit",
                          filename=attachment.filename,
                          size_mb=attachment.content_length / (1024 * 1024),
                          max_size_mb=max_size_mb)
            return False
        
        return True
    
    async def extract_text(self, attachment: AttachmentInfo) -> Optional[str]:
        """Extract text from attachment if possible."""
        try:
            mimetype = attachment.content_type
            
            if mimetype in self.supported_mimetypes:
                extractor = self.supported_mimetypes[mimetype]
                text = await extractor(attachment)
                
                if text:
                    logger.info("Text extracted from attachment",
                               filename=attachment.filename,
                               mimetype=mimetype,
                               text_length=len(text))
                    return text
                else:
                    logger.warning("No text extracted from attachment",
                                 filename=attachment.filename,
                                 mimetype=mimetype)
                    return None
            else:
                logger.info("Mimetype not supported for text extraction",
                           filename=attachment.filename,
                           mimetype=mimetype)
                return None
                
        except Exception as exc:
            logger.error("Failed to extract text from attachment",
                        filename=attachment.filename,
                        mimetype=attachment.content_type,
                        exc_info=exc)
            return None
    
    async def _extract_pdf_text(self, attachment: AttachmentInfo) -> Optional[str]:
        """Extract text from PDF attachment."""
        try:
            # Try to extract text using PyMuPDF
            import fitz  # PyMuPDF
            
            # Load PDF from memory
            pdf_document = fitz.open(stream=attachment.content, filetype="pdf")
            
            text_content = []
            for page_num in range(pdf_document.page_count):
                page = pdf_document.load_page(page_num)
                text_content.append(page.get_text())
            
            pdf_document.close()
            
            extracted_text = '\n'.join(text_content)
            
            # Check if we got meaningful text
            if len(extracted_text.strip()) > 50:  # Minimum text threshold
                return extracted_text
            else:
                logger.info("PDF appears to be image-based, no text extracted",
                           filename=attachment.filename)
                return None
                
        except ImportError:
            logger.warning("PyMuPDF not available for PDF text extraction")
            return None
        except Exception as exc:
            logger.error("PDF text extraction failed",
                        filename=attachment.filename,
                        exc_info=exc)
            return None
    
    async def _extract_doc_text(self, attachment: AttachmentInfo) -> Optional[str]:
        """Extract text from DOC attachment."""
        try:
            # Try to extract text using python-docx (may not work for .doc files)
            from docx import Document
            
            # Load document from memory
            doc = Document(attachment.content)
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            extracted_text = '\n'.join(text_content)
            
            if extracted_text.strip():
                return extracted_text
            else:
                logger.warning("No text extracted from DOC file",
                             filename=attachment.filename)
                return None
                
        except ImportError:
            logger.warning("python-docx not available for DOC text extraction")
            return None
        except Exception as exc:
            logger.error("DOC text extraction failed",
                        filename=attachment.filename,
                        exc_info=exc)
            return None
    
    async def _extract_docx_text(self, attachment: AttachmentInfo) -> Optional[str]:
        """Extract text from DOCX attachment."""
        try:
            from docx import Document
            
            # Load document from memory
            doc = Document(attachment.content)
            
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            extracted_text = '\n'.join(text_content)
            
            if extracted_text.strip():
                return extracted_text
            else:
                logger.warning("No text extracted from DOCX file",
                             filename=attachment.filename)
                return None
                
        except ImportError:
            logger.warning("python-docx not available for DOCX text extraction")
            return None
        except Exception as exc:
            logger.error("DOCX text extraction failed",
                        filename=attachment.filename,
                        exc_info=exc)
            return None
    
    async def _extract_excel_text(self, attachment: AttachmentInfo) -> Optional[str]:
        """Extract text from Excel attachment."""
        try:
            # Try to extract text using openpyxl
            from openpyxl import load_workbook
            
            # Load workbook from memory
            workbook = load_workbook(attachment.content, read_only=True)
            
            text_content = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_content.append(f"Sheet: {sheet_name}")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell_value in row:
                        if cell_value is not None:
                            row_text.append(str(cell_value))
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            workbook.close()
            
            extracted_text = '\n'.join(text_content)
            
            if extracted_text.strip():
                return extracted_text
            else:
                logger.warning("No text extracted from Excel file",
                             filename=attachment.filename)
                return None
                
        except ImportError:
            logger.warning("openpyxl not available for Excel text extraction")
            return None
        except Exception as exc:
            logger.error("Excel text extraction failed",
                        filename=attachment.filename,
                        exc_info=exc)
            return None
    
    async def _extract_xlsx_text(self, attachment: AttachmentInfo) -> Optional[str]:
        """Extract text from XLSX attachment."""
        try:
            from openpyxl import load_workbook
            
            # Load workbook from memory
            workbook = load_workbook(attachment.content, read_only=True)
            
            text_content = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_content.append(f"Sheet: {sheet_name}")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell_value in row:
                        if cell_value is not None:
                            row_text.append(str(cell_value))
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            workbook.close()
            
            extracted_text = '\n'.join(text_content)
            
            if extracted_text.strip():
                return extracted_text
            else:
                logger.warning("No text extracted from XLSX file",
                             filename=attachment.filename)
                return None
                
        except ImportError:
            logger.warning("openpyxl not available for XLSX text extraction")
            return None
        except Exception as exc:
            logger.error("XLSX text extraction failed",
                        filename=attachment.filename,
                        exc_info=exc)
            return None
    
    async def _extract_plain_text(self, attachment: AttachmentInfo) -> Optional[str]:
        """Extract text from plain text attachment."""
        try:
            # Try to decode the content
            if isinstance(attachment.content, bytes):
                # Try different encodings
                encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
                
                for encoding in encodings:
                    try:
                        text = attachment.content.decode(encoding)
                        if text.strip():
                            return text
                    except UnicodeDecodeError:
                        continue
                
                # If all encodings fail, use errors='ignore'
                text = attachment.content.decode('utf-8', errors='ignore')
                return text if text.strip() else None
            else:
                return str(attachment.content) if attachment.content else None
                
        except Exception as exc:
            logger.error("Plain text extraction failed",
                        filename=attachment.filename,
                        exc_info=exc)
            return None
    
    async def _extract_html_text(self, attachment: AttachmentInfo) -> Optional[str]:
        """Extract text from HTML attachment."""
        try:
            from bs4 import BeautifulSoup
            
            # Decode HTML content
            if isinstance(attachment.content, bytes):
                html_content = attachment.content.decode('utf-8', errors='ignore')
            else:
                html_content = str(attachment.content)
            
            # Parse HTML and extract text
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text content
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            if text.strip():
                return text
            else:
                logger.warning("No text extracted from HTML file",
                             filename=attachment.filename)
                return None
                
        except ImportError:
            logger.warning("BeautifulSoup not available for HTML text extraction")
            return None
        except Exception as exc:
            logger.error("HTML text extraction failed",
                        filename=attachment.filename,
                        exc_info=exc)
            return None


class OCRTaskManager:
    """Manages OCR tasks for image and PDF attachments."""
    
    def __init__(self):
        """Initialize OCR task manager."""
        self.pending_tasks = []
        self.completed_tasks = []
        self.failed_tasks = []
    
    async def create_ocr_task(self, attachment: AttachmentInfo, email_id: str) -> OCRTask:
        """Create a new OCR task."""
        try:
            # Generate task ID
            task_id = f"ocr_{hash(attachment.content_hash)}_{hash(attachment.filename)}"
            
            # Create OCR task
            ocr_task = OCRTask(
                task_id=task_id,
                email_id=email_id,
                attachment_id=attachment.filename,
                content_type=attachment.content_type,
                content_hash=attachment.content_hash,
                content_length=attachment.content_length
            )
            
            # Add to pending tasks
            self.pending_tasks.append(ocr_task)
            
            logger.info("OCR task created",
                       task_id=task_id,
                       filename=attachment.filename,
                       content_type=attachment.content_type)
            
            return ocr_task
            
        except Exception as exc:
            logger.error("Failed to create OCR task",
                        filename=attachment.filename,
                        exc_info=exc)
            raise
    
    async def get_pending_tasks(self, limit: int = 100) -> list[OCRTask]:
        """Get pending OCR tasks."""
        return self.pending_tasks[:limit]
    
    async def mark_task_completed(self, task_id: str, extracted_text: str, confidence: float = None):
        """Mark OCR task as completed."""
        try:
            # Find the task
            task = next((t for t in self.pending_tasks if t.task_id == task_id), None)
            
            if task:
                # Update task
                task.status = ProcessingStatus.COMPLETED
                task.extracted_text = extracted_text
                task.extracted_text_hash = self._calculate_hash(extracted_text.encode('utf-8'))
                task.confidence_score = confidence
                task.completed_at = datetime.utcnow()
                
                # Move to completed tasks
                self.pending_tasks.remove(task)
                self.completed_tasks.append(task)
                
                logger.info("OCR task completed",
                           task_id=task_id,
                           text_length=len(extracted_text),
                           confidence=confidence)
            else:
                logger.warning("OCR task not found", task_id=task_id)
                
        except Exception as exc:
            logger.error("Failed to mark OCR task completed",
                        task_id=task_id,
                        exc_info=exc)
    
    async def mark_task_failed(self, task_id: str, error_type: ErrorType, error_message: str):
        """Mark OCR task as failed."""
        try:
            # Find the task
            task = next((t for t in self.pending_tasks if t.task_id == task_id), None)
            
            if task:
                # Update task
                task.status = ProcessingStatus.FAILED
                task.error_type = error_type
                task.error_message = error_message
                
                # Move to failed tasks
                self.pending_tasks.remove(task)
                self.failed_tasks.append(task)
                
                logger.warning("OCR task failed",
                              task_id=task_id,
                              error_type=error_type,
                              error_message=error_message)
            else:
                logger.warning("OCR task not found", task_id=task_id)
                
        except Exception as exc:
            logger.error("Failed to mark OCR task failed",
                        task_id=task_id,
                        exc_info=exc)
    
    def _calculate_hash(self, content: bytes) -> str:
        """Calculate SHA-256 hash of content."""
        return hashlib.sha256(content).hexdigest()
    
    def get_task_stats(self) -> Dict[str, int]:
        """Get OCR task statistics."""
        return {
            'pending': len(self.pending_tasks),
            'completed': len(self.completed_tasks),
            'failed': len(self.failed_tasks),
            'total': len(self.pending_tasks) + len(self.completed_tasks) + len(self.failed_tasks)
        }
